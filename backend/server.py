"""
EcoRevive Backend API Server
============================
FastAPI server that handles the full pipeline:
1. Download Sentinel-2 imagery from Earth Engine
2. Run California Fire Model inference
3. Generate Gemini reasoning analysis
4. Return results to frontend
"""

import os
import sys
import base64
import json
import io
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime
import ee
from google.oauth2 import service_account

import numpy as np
from PIL import Image
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

# Add project paths
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))
sys.path.insert(0, str(PROJECT_ROOT / "California-Fire-Model"))

# Load environment variables
load_dotenv(PROJECT_ROOT / ".env")

# Import our modules
from ee_download import initialize_ee, download_sentinel2_for_model, create_rgb_from_bands

# Initialize FastAPI app
app = FastAPI(
    title="EcoRevive API",
    description="Burn severity analysis using satellite imagery and AI",
    version="1.0.0"
)

# Enable CORS for frontend (local + production)
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "http://localhost:3000",
        "http://127.0.0.1:5173",
        "https://ecorevive.vercel.app",
        "https://ecorevive-original.vercel.app",
        "*",  # Allow all origins for hackathon demo
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Global state
ee_initialized = False
model = None
device = None


# Request/Response Models
class AnalyzeRequest(BaseModel):
    west: float
    south: float
    east: float
    north: float
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    user_type: Optional[str] = "personal"


class AnalyzeResponse(BaseModel):
    success: bool
    satellite_image: Optional[str] = None  # Base64 encoded RGB satellite image
    severity_image: Optional[str] = None  # Base64 encoded colorized PNG
    raw_severity_image: Optional[str] = None  # Base64 encoded raw grayscale model output
    severity_stats: Optional[Dict[str, Any]] = None
    gemini_analysis: Optional[str] = None  # Legacy verbose text for display
    layer2_output: Optional[Dict[str, Any]] = None  # Structured Layer 2 JSON
    layer3_context: Optional[Dict[str, Any]] = None  # Layer 3 contextual analysis (urban detection, cautions)
    carbon_analysis: Optional[Dict[str, Any]] = None  # Carbon sequestration calculator
    fire_verified: bool = True  # Whether actual fire activity was detected
    fire_verification_message: Optional[str] = None  # Message if fire not verified
    error: Optional[str] = None


class Layer2AnalyzeResponse(BaseModel):
    """Structured Layer 2 output for Layer 3 consumption."""
    success: bool
    layer2_output: Optional[Dict[str, Any]] = None  # Full structured Layer 2 data
    satellite_image: Optional[str] = None  # Base64 encoded RGB for reference
    severity_image: Optional[str] = None  # Base64 encoded colorized severity
    error: Optional[str] = None


class ChatRequest(BaseModel):
    """Request for the AI chat endpoint."""
    message: str  # User's message or quick action prompt
    action_type: Optional[str] = None  # Quick action ID (safety, hope, species, etc.)
    user_type: str = "personal"  # personal or professional
    context: Optional[Dict[str, Any]] = None  # Analysis context (layer2_output, severity_stats, bbox)


class ChatResponse(BaseModel):
    """Response from the AI chat endpoint."""
    success: bool
    response: Optional[str] = None
    error: Optional[str] = None


class PDFExportRequest(BaseModel):
    """Request for PDF export endpoint."""
    satellite_image: str
    severity_image: str
    severity_stats: Dict[str, Any]
    bbox: Dict[str, float]
    layer2_output: Optional[Dict[str, Any]] = None
    layer3_context: Optional[Dict[str, Any]] = None
    carbon_analysis: Optional[Dict[str, Any]] = None
    report_type: str = "personal"
    user_type: str = "personal"
    location_name: Optional[str] = None
    analysis_id: Optional[str] = None
    chat_history: Optional[List[Dict[str, Any]]] = None  # Chat messages to append to report


class PDFExportResponse(BaseModel):
    """Response from PDF export endpoint."""
    success: bool
    pdf_base64: Optional[str] = None
    filename: Optional[str] = None
    error: Optional[str] = None


class WordExportRequest(BaseModel):
    """Request for Word export endpoint."""
    satellite_image: str
    severity_image: str
    severity_stats: Dict[str, Any]
    bbox: Dict[str, float]
    layer2_output: Optional[Dict[str, Any]] = None
    layer3_context: Optional[Dict[str, Any]] = None
    carbon_analysis: Optional[Dict[str, Any]] = None
    report_type: str = "personal"
    user_type: str = "personal"
    location_name: Optional[str] = None
    chat_history: Optional[List[Dict[str, Any]]] = None  # Chat messages to append to report


class WordExportResponse(BaseModel):
    """Response from Word export endpoint."""
    success: bool
    docx_base64: Optional[str] = None
    filename: Optional[str] = None
    error: Optional[str] = None


class HopeVisualizationRequest(BaseModel):
    """Request for hope visualization endpoint."""
    ecosystem_type: str = "mixed_conifer"
    years_in_future: int = 15
    mean_severity: float = 0.5
    area_hectares: float = 100.0
    bbox: Optional[Dict[str, float]] = None


class HopeVisualizationResponse(BaseModel):
    """Response from hope visualization endpoint."""
    success: bool
    image_base64: Optional[str] = None
    forecast: Optional[Dict[str, Any]] = None
    description: Optional[str] = None
    error: Optional[str] = None


def create_synthetic_image(bbox: Dict[str, float]) -> np.ndarray:
    """
    Create synthetic Sentinel-2 imagery for testing when EE is unavailable.
    Returns a (10, 256, 256) array simulating multispectral data.
    """
    # Create base noise patterns
    np.random.seed(42)  # Reproducible for demos
    
    # Create spatially coherent patterns using simple gradients + noise
    h, w = 256, 256
    
    # Base terrain pattern
    x = np.linspace(0, 4 * np.pi, w)
    y = np.linspace(0, 4 * np.pi, h)
    xx, yy = np.meshgrid(x, y)
    terrain = np.sin(xx) * np.cos(yy) * 0.3 + 0.5
    
    # Add some random burn patterns
    burn_center_x = np.random.randint(50, 200)
    burn_center_y = np.random.randint(50, 200)
    distances = np.sqrt((np.arange(w) - burn_center_x)**2 + (np.arange(h)[:, None] - burn_center_y)**2)
    burn_pattern = np.exp(-distances / 40) * 0.6
    
    # Create 10 bands with realistic spectral relationships
    bands = []
    for i in range(10):
        # Each band has slightly different characteristics
        noise = np.random.randn(h, w) * 0.1
        band = terrain + burn_pattern * (0.5 + i * 0.05) + noise
        
        # Scale to typical Sentinel-2 reflectance values (0-10000)
        band = (band * 3000 + 1000).clip(0, 10000)
        bands.append(band.astype(np.float32))
    
    return np.stack(bands, axis=0)


def load_fire_model():
    """Load the California Fire Model."""
    global model, device
    
    try:
        import torch
        from model.architecture import CaliforniaFireModel
        
        checkpoint_path = PROJECT_ROOT / "California-Fire-Model/checkpoints/model.pth"
        
        if not checkpoint_path.exists():
            print(f"[WARNING] Model checkpoint not found at {checkpoint_path}")
            return False
            
        device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
        print(f"[INFO] Loading Fire Model on {device}...")
        
        model = CaliforniaFireModel(
            input_channels=10,
            output_channels=1,
            base_channels=64,
            use_attention=True,
        )
        
        checkpoint = torch.load(checkpoint_path, map_location=device, weights_only=False)
        
        if isinstance(checkpoint, dict) and 'model_state_dict' in checkpoint:
            model.load_state_dict(checkpoint['model_state_dict'])
        elif isinstance(checkpoint, dict) and 'state_dict' in checkpoint:
            model.load_state_dict(checkpoint['state_dict'])
        else:
            model.load_state_dict(checkpoint)
            
        model.to(device)
        model.eval()
        
        print(f"[OK] Fire Model loaded ({sum(p.numel() for p in model.parameters()):,} parameters)")
        return True
        
    except Exception as e:
        print(f"[ERROR] Failed to load Fire Model: {e}")
        import traceback
        traceback.print_exc()
        return False


def run_inference(image_array: np.ndarray) -> tuple:
    """
    Run Fire Model inference on image array.
    Uses EXACT normalization from training.
    
    Args:
        image_array: (10, 256, 256) Sentinel-2 bands
        
    Returns:
        Tuple of (severity_map, stats_dict)
    """
    global model, device
    
    import torch
    from config import get_band_stats
    
    # Get band statistics
    means, stds = get_band_stats()
    means = np.array(means).reshape(-1, 1, 1)
    stds = np.array(stds).reshape(-1, 1, 1)
    
    # EXACT normalization from training:
    # 1. Replace NaN with 0
    image_array = np.nan_to_num(image_array, nan=0.0)
    
    # 2. Z-score normalize
    normalized = (image_array - means) / (stds + 1e-6)
    
    # 3. Clip to [-3, 3]
    normalized = np.clip(normalized, -3, 3)
    
    # 4. Scale to [0, 1] via (x + 3) / 6
    normalized = (normalized + 3) / 6
    
    # Convert to tensor
    tensor = torch.from_numpy(normalized).float().unsqueeze(0).to(device)
    
    # Run inference
    with torch.no_grad():
        output = model(tensor)
        severity_map = torch.sigmoid(output).squeeze().cpu().numpy()
    
    # Compute statistics
    stats = {
        'mean_severity': float(severity_map.mean()),
        'max_severity': float(severity_map.max()),
        'min_severity': float(severity_map.min()),
        'burned_ratio': float((severity_map > 0.5).mean()),
        'high_severity_ratio': float((severity_map > 0.75).mean()),
        'moderate_severity_ratio': float(((severity_map > 0.25) & (severity_map <= 0.75)).mean()),
        'low_severity_ratio': float((severity_map <= 0.25).mean()),
    }
    
    return severity_map, stats


def run_tiled_inference(tiles: list, metadata: dict) -> tuple:
    """
    Run inference on multiple tiles and stitch results together.
    
    Args:
        tiles: List of tile dicts from download_for_inference
        metadata: Metadata dict with n_rows, n_cols
        
    Returns:
        Tuple of (stitched_severity_map, stitched_satellite_image, combined_stats)
    """
    n_rows = metadata['n_rows']
    n_cols = metadata['n_cols']
    tile_size = 256
    
    # Create output arrays
    stitched_severity = np.zeros((n_rows * tile_size, n_cols * tile_size), dtype=np.float32)
    stitched_satellite = np.zeros((n_rows * tile_size, n_cols * tile_size, 3), dtype=np.uint8)
    
    all_severities = []
    
    print(f"   Processing {len(tiles)} tiles...")
    
    for tile in tiles:
        row = tile['row']
        col = tile['col']
        image = tile['image']
        
        # Run inference on this tile
        severity_map, _ = run_inference(image)
        all_severities.append(severity_map)
        
        # Create RGB for satellite view
        rgb = create_tile_rgb(image)
        
        # Calculate position in stitched image
        y_start = (n_rows - 1 - row) * tile_size  # Flip rows (row 0 is at bottom)
        y_end = y_start + tile_size
        x_start = col * tile_size
        x_end = x_start + tile_size
        
        # Place in stitched arrays
        stitched_severity[y_start:y_end, x_start:x_end] = severity_map
        stitched_satellite[y_start:y_end, x_start:x_end] = rgb
    
    # Compute combined statistics
    if all_severities:
        combined = np.concatenate([s.flatten() for s in all_severities])
        stats = {
            'mean_severity': float(combined.mean()),
            'max_severity': float(combined.max()),
            'min_severity': float(combined.min()),
            'burned_ratio': float((combined > 0.5).mean()),
            'high_severity_ratio': float((combined > 0.75).mean()),
            'moderate_severity_ratio': float(((combined > 0.25) & (combined <= 0.75)).mean()),
            'low_severity_ratio': float((combined <= 0.25).mean()),
            'n_tiles_processed': len(tiles),
        }
    else:
        stats = {'mean_severity': 0, 'n_tiles_processed': 0}
    
    print(f"   Stitched output: {stitched_severity.shape}")
    
    return stitched_severity, stitched_satellite, stats


def create_tile_rgb(image_array: np.ndarray) -> np.ndarray:
    """Create RGB array from Sentinel-2 bands for a single tile."""
    # Band statistics from training
    BAND_MEANS = [472.8, 673.8, 770.8, 1087.8, 1747.6, 1997.1, 2106.4, 2188.9, 1976.1, 1404.5]
    BAND_STDS = [223.7, 255.4, 345.6, 313.5, 366.7, 417.6, 476.7, 437.1, 472.7, 438.4]
    
    # Normalize bands
    normalized = np.zeros_like(image_array, dtype=np.float32)
    for i in range(10):
        normalized[i] = (image_array[i] - BAND_MEANS[i]) / (BAND_STDS[i] + 1e-6)
    normalized = np.clip(normalized, -3, 3)
    normalized = (normalized + 3) / 6
    
    # False color B5, B4, B3
    rgb = np.stack([normalized[3], normalized[2], normalized[1]], axis=-1)
    rgb = np.clip(rgb, 0, 1)
    rgb = (rgb * 255).astype(np.uint8)
    
    return rgb


def severity_to_image(severity_map: np.ndarray) -> str:
    """Convert severity map to colorized PNG using 'hot' colormap (matching training visualization)."""
    from PIL import Image
    import matplotlib.pyplot as plt
    import matplotlib.cm as cm
    
    # Use matplotlib's 'hot' colormap (black -> red -> yellow -> white)
    # This EXACTLY matches the training visualization
    cmap = cm.get_cmap('hot')
    
    # Apply colormap: severity (0-1) -> RGBA -> RGB
    rgba = cmap(severity_map)  # Returns (H, W, 4)
    rgb = (rgba[:, :, :3] * 255).astype(np.uint8)
    
    # Create image
    img = Image.fromarray(rgb)
    
    # Encode to base64
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    base64_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
    
    return f"data:image/png;base64,{base64_str}"


def rgb_array_to_base64(rgb_array: np.ndarray) -> str:
    """Convert RGB numpy array to base64 PNG string."""
    from PIL import Image
    
    img = Image.fromarray(rgb_array, mode='RGB')
    
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    base64_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
    
    return f"data:image/png;base64,{base64_str}"


def severity_to_raw_image(severity_map: np.ndarray) -> str:
    """Convert severity map to raw grayscale PNG base64 string (actual model output)."""
    from PIL import Image
    
    # Convert to 8-bit grayscale (0=no burn, 255=max burn)
    grayscale = (severity_map * 255).astype(np.uint8)
    
    # Create grayscale image
    img = Image.fromarray(grayscale, mode='L')
    
    # Encode to base64
    buffer = io.BytesIO()
    img.save(buffer, format='PNG')
    base64_str = base64.b64encode(buffer.getvalue()).decode('utf-8')
    
    return f"data:image/png;base64,{base64_str}"


def get_gemini_analysis(
    stats: Dict[str, Any], 
    bbox: Dict[str, float], 
    user_type: str,
    satellite_rgb: np.ndarray = None,
    severity_map: np.ndarray = None
) -> Dict[str, Any]:
    """
    Get TRUE MULTIMODAL Gemini analysis of the burn severity results.
    
    This function sends BOTH the satellite image AND severity overlay to Gemini,
    enabling spatial reasoning about WHAT is burned, not just statistics.
    
    Args:
        stats: Burn severity statistics dict
        bbox: Bounding box dict with west, south, east, north
        user_type: "professional" or "personal"
        satellite_rgb: (H, W, 3) RGB satellite image array
        severity_map: (H, W) severity predictions from U-Net
        
    Returns:
        Dict with 'text' (formatted analysis) and 'analysis' (structured data)
    """
    try:
        # Check if we have images for true multimodal analysis
        if satellite_rgb is not None and severity_map is not None:
            return _get_multimodal_analysis(stats, bbox, user_type, satellite_rgb, severity_map)
        else:
            # Fallback to text-only analysis
            return _get_text_only_analysis(stats, bbox, user_type)
            
    except Exception as e:
        print(f"[WARNING] Gemini analysis failed: {e}")
        import traceback
        traceback.print_exc()
        fallback_text = (
            f"AI analysis temporarily unavailable. Manual assessment: "
            f"{stats['mean_severity']:.0%} average severity with "
            f"{stats['high_severity_ratio']:.0%} high-severity areas requiring attention."
        )
        return {'text': fallback_text, 'analysis': None}


def _get_multimodal_analysis(
    stats: Dict[str, Any],
    bbox: Dict[str, float],
    user_type: str,
    satellite_rgb: np.ndarray,
    severity_map: np.ndarray
) -> Dict[str, Any]:
    """
    TRUE MULTIMODAL analysis - sends images to Gemini for spatial reasoning.
    """
    from reasoning.gemini_multimodal import (
        MultimodalAnalyzer,
        create_image_pack,
        build_gemini_context
    )
    from reasoning import create_client
    
    print("   [INFO] Running TRUE MULTIMODAL analysis (sending images to Gemini)...")
    
    # Create Gemini client
    client = create_client()
    
    # Calculate center location
    center_lat = (bbox['north'] + bbox['south']) / 2
    center_lon = (bbox['west'] + bbox['east']) / 2
    location = (center_lat, center_lon)
    
    # Prepare RGB tile for multimodal - transpose from (H, W, 3) to (3, H, W)
    if satellite_rgb.ndim == 3 and satellite_rgb.shape[2] == 3:
        rgb_tile = np.moveaxis(satellite_rgb, -1, 0)
    else:
        rgb_tile = satellite_rgb
    
    # Build metadata for context
    metadata = {
        'bbox': bbox,
        'user_type': user_type,
        'fire_date': None,  # Could be passed from request
        'days_since_fire': None,
    }
    
    # Run multimodal analysis
    analyzer = MultimodalAnalyzer(client)
    result = analyzer.analyze(
        rgb_tile=rgb_tile,
        severity_map=severity_map,
        location=location,
        metadata=metadata,
        unet_confidence=0.85  # Assumed confidence
    )
    
    # Format the response for the frontend
    if result.get('status') == 'complete':
        analysis = result.get('analysis', {})
        
        # Build human-readable text from spatial analysis
        text_parts = []
        
        # Visual grounding section
        vg = analysis.get('visual_grounding', {})
        if vg:
            land_cover = ", ".join(vg.get('observed_land_cover', ['Unknown']))
            text_parts.append(f"## What Gemini Sees in the Satellite Image\n")
            text_parts.append(f"**Land Cover Types Detected:** {land_cover}\n")
            if vg.get('terrain_features'):
                text_parts.append(f"**Terrain Features:** {', '.join(vg.get('terrain_features', []))}\n")
            if vg.get('pre_fire_vegetation_description'):
                text_parts.append(f"**Vegetation Description:** {vg.get('pre_fire_vegetation_description')}\n")
        
        # Segmentation quality section
        sq = analysis.get('segmentation_quality', {})
        if sq:
            quality = sq.get('overall_quality', 'unknown')
            confidence = sq.get('confidence_in_prediction', 0)
            text_parts.append(f"\n## U-Net Prediction Quality Assessment\n")
            text_parts.append(f"**Overall Quality:** {quality.upper()}\n")
            text_parts.append(f"**Gemini's Confidence in Predictions:** {confidence:.0%}\n")
            if sq.get('artifact_flags'):
                text_parts.append(f"**[WARNING] Artifacts Detected:** {', '.join(sq.get('artifact_flags'))}\n")
            if sq.get('quality_notes'):
                text_parts.append(f"**Notes:** {sq.get('quality_notes')}\n")
        
        # Spatial patterns section
        sp = analysis.get('spatial_patterns', {})
        if sp:
            text_parts.append(f"\n## Spatial Pattern Analysis\n")
            
            frag = sp.get('fragmentation_assessment', {})
            if frag:
                text_parts.append(f"**Burn Patches Visible:** {frag.get('patch_count_visual', 'unknown')}\n")
                text_parts.append(f"**Connectivity:** {frag.get('connectivity', 'unknown')}\n")
                text_parts.append(f"**Shape Complexity:** {frag.get('shape_complexity', 'unknown')}\n")
            
            edge = sp.get('edge_characteristics', {})
            if edge:
                text_parts.append(f"**Edge Sharpness:** {edge.get('edge_sharpness', 'unknown')}\n")
                if edge.get('unburned_inclusions'):
                    text_parts.append(f"**Unburned Islands:** Yes - {edge.get('inclusion_significance', '')}\n")
            
            grad = sp.get('gradient_analysis', {})
            if grad:
                if grad.get('dominant_direction') and grad.get('dominant_direction') != 'none':
                    text_parts.append(f"**Fire Spread Direction:** {grad.get('dominant_direction')}\n")
                if grad.get('fire_behavior_inference'):
                    text_parts.append(f"**Fire Behavior Inference:** {grad.get('fire_behavior_inference')}\n")
        
        # Ecological interpretation
        ei = analysis.get('ecological_interpretation', {})
        if ei:
            text_parts.append(f"\n## Ecological Interpretation\n")
            text_parts.append(f"**Natural Regeneration Potential:** {ei.get('natural_regeneration_potential', 'unknown')}\n")
            text_parts.append(f"**Seed Source Availability:** {ei.get('seed_source_availability', 'unknown')}\n")
            if ei.get('regeneration_rationale'):
                text_parts.append(f"**Rationale:** {ei.get('regeneration_rationale')}\n")
            
            # Differential impacts
            impacts = ei.get('differential_impacts', [])
            if impacts:
                text_parts.append(f"\n**Differential Impacts:**\n")
                for impact in impacts[:3]:
                    text_parts.append(f"- {impact.get('vegetation_type', 'Unknown')}: "
                                     f"{impact.get('severity_level', '')} severity - "
                                     f"{impact.get('ecological_significance', '')}\n")
        
        # Priority zones section
        zones = analysis.get('priority_zones', [])
        if zones:
            text_parts.append(f"\n## Priority Restoration Zones\n")
            for zone in zones[:5]:
                urgency_label = {
                    'immediate': '[URGENT]',
                    '6_months': '[6 MONTHS]',
                    '1_year': '[1 YEAR]',
                    '2_3_years': '[2-3 YEARS]'
                }.get(zone.get('urgency', ''), '[--]')

                text_parts.append(f"\n### {urgency_label} Zone {zone.get('zone_id', '?')}: {zone.get('urgency', 'unknown').replace('_', ' ').title()}\n")
                text_parts.append(f"**Location:** {zone.get('location_description', 'See overlay')}\n")
                text_parts.append(f"**Severity:** {zone.get('severity', 'unknown')}\n")
                text_parts.append(f"**Why Priority:** {zone.get('priority_reason', 'N/A')}\n")
                text_parts.append(f"**Recommended Action:** {zone.get('recommended_intervention', 'N/A').replace('_', ' ').title()}\n")
        
        # Machine-readable signals section
        signals = analysis.get('signals_for_final_model', {})
        if signals:
            text_parts.append(f"\n## Restoration Scores (Machine-Readable)\n")
            text_parts.append(f"| Metric | Score |\n|--------|-------|\n")
            text_parts.append(f"| Restoration Potential | {signals.get('restoration_potential_score', 0):.2f} |\n")
            text_parts.append(f"| Intervention Urgency | {signals.get('intervention_urgency_score', 0):.2f} |\n")
            text_parts.append(f"| Ecological Complexity | {signals.get('ecological_complexity_score', 0):.2f} |\n")
            text_parts.append(f"| Risk Score | {signals.get('risk_score', 0):.2f} |\n")
        
        # Reasoning trace
        if analysis.get('reasoning_trace'):
            text_parts.append(f"\n---\n*Gemini's reasoning: {analysis.get('reasoning_trace')}*\n")
        
        formatted_text = "".join(text_parts)
        
        # Add human review warning if needed
        if result.get('human_review', {}).get('required'):
            triggers = result.get('human_review', {}).get('triggers', [])
            formatted_text = (
                f"**[WARNING] Human Review Recommended:** {', '.join(triggers)}\n\n"
                + formatted_text
            )
        
        print(f"   [OK] Multimodal analysis complete! Found {len(zones)} priority zones.")
        
        return {
            'text': formatted_text,
            'analysis': analysis,
            'signals': signals,
            'human_review_required': result.get('human_review', {}).get('required', False)
        }
    
    else:
        # Analysis failed - return error info
        error_text = f"[WARNING] Multimodal analysis failed: {result.get('error', result.get('status', 'unknown'))}"
        print(f"   {error_text}")
        return {'text': error_text, 'analysis': None}


def _get_text_only_analysis(stats: Dict[str, Any], bbox: Dict[str, float], user_type: str) -> Dict[str, Any]:
    """
    LEGACY text-only analysis (fallback when images aren't available).
    """
    from reasoning import EcoReviveGemini
    
    print("   [WARNING] Running LEGACY text-only analysis (no images sent to Gemini)")
    
    client = EcoReviveGemini()
    
    # Format prompt based on user type
    if user_type == "professional":
        prompt = f"""You are an expert wildfire restoration ecologist. Analyze this burn severity assessment:

**Location**: {bbox['south']:.4f}°N to {bbox['north']:.4f}°N, {bbox['west']:.4f}°W to {bbox['east']:.4f}°W

**Burn Severity Analysis**:
- Mean Severity: {stats['mean_severity']:.1%}
- Maximum Severity: {stats['max_severity']:.1%}
- High Severity Area (>75%): {stats['high_severity_ratio']:.1%}
- Moderate Severity Area (25-75%): {stats['moderate_severity_ratio']:.1%}
- Low/Unburned Area (<25%): {stats['low_severity_ratio']:.1%}

Provide a professional assessment including:
1. **Severity Classification**: Overall burn severity category
2. **Ecological Impact**: Expected vegetation and habitat damage
3. **Restoration Priority**: Urgency and recommended approach
4. **Key Species to Plant**: Native species suited for this recovery
5. **Timeline**: Expected natural vs assisted recovery time
6. **Budget Estimate**: Rough cost per hectare for restoration

Keep the response concise but actionable for grant proposals and project planning."""
    else:
        prompt = f"""You are a friendly environmental guide helping someone understand wildfire damage. Explain this burn severity assessment in simple terms:

**Location**: Near {bbox['south']:.2f}°N, {abs(bbox['west']):.2f}°W

**What the satellite analysis found**:
- Overall burn damage: {stats['mean_severity']:.0%}
- Severely burned areas: {stats['high_severity_ratio']:.0%} of the region
- Moderately burned: {stats['moderate_severity_ratio']:.0%}
- Lightly affected or untouched: {stats['low_severity_ratio']:.0%}

Please explain:
1. What does this mean for the land and wildlife?
2. How long will it take nature to recover?
3. What can volunteers do to help?
4. What plants and trees will grow back first?

Keep it encouraging and accessible - this person cares about the environment but isn't a scientist."""

    response = client.analyze_multimodal(prompt=prompt, use_json=False)
    return {'text': response.get('text', 'Analysis not available.'), 'analysis': None}

###TESTING For google cloud deployment###
def initialize_ee_robust():
    """Robust Earth Engine initialization for Cloud Run using Service Account."""
    try:
        # 1. Look for the secret file we mounted in the Docker command
        # Defaults to /secrets/service-account.json if env var isn't set
        creds_path = os.environ.get('GOOGLE_APPLICATION_CREDENTIALS', '/secrets/service-account.json')
        
        if os.path.exists(creds_path):
            print(f"[INFO] Found Service Account at {creds_path}")
            
            # Load credentials with the specific Earth Engine scope
            credentials = service_account.Credentials.from_service_account_file(creds_path)
            scoped_credentials = credentials.with_scopes([
                'https://www.googleapis.com/auth/earthengine',
                'https://www.googleapis.com/auth/cloud-platform'
            ])
            
            # Initialize with these scoped credentials
            ee.Initialize(scoped_credentials)
            print("[SUCCESS] EE Initialized via Service Account!")
            return True
            
        else:
            print(f"[WARNING] Service account key not found at {creds_path}")
            
        # 2. Fallback for Local Development (uses local gcloud auth if available)
        print("[INFO] Attempting default local authentication...")
        ee.Initialize()
        return True
        
    except Exception as e:
        print(f"[ERROR] EE failed to initialize: {e}")
        return False





@app.on_event("startup")
async def startup_event():
    """Initialize services on startup."""
    global ee_initialized
    
    print("[INFO] Starting EcoRevive API Server...")
    
    # # Check for encoded credentials (e.g. from Railway/Docker env)
    # if 'GOOGLE_APPLICATION_CREDENTIALS_JSON' in os.environ:
    #     try:
    #         creds_json = base64.b64decode(os.environ['GOOGLE_APPLICATION_CREDENTIALS_JSON']).decode('utf-8')
    #         print("[INFO] Found GOOGLE_APPLICATION_CREDENTIALS_JSON, decoding...")
    #         # Pass to ee_download via its expected env var
    #         os.environ['EE_SERVICE_ACCOUNT_JSON'] = creds_json
    #     except Exception as e:
    #         print(f"[ERROR] Failed to decode GOOGLE_APPLICATION_CREDENTIALS_JSON: {e}")
    
    
    # Initialize Earth Engine using the new robust function
    ee_initialized = initialize_ee_robust()
    
    # Load Fire Model
    load_fire_model()
    
    print("[OK] Server ready!")


@app.get("/")
async def root():
    return {
        "service": "EcoRevive API",
        "status": "running",
        "ee_initialized": ee_initialized,
        "model_loaded": model is not None
    }


@app.get("/api/search")
async def search_location(q: str):
    """Proxy Nominatim search to avoid browser CORS issues."""
    import requests as req
    resp = req.get(
        "https://nominatim.openstreetmap.org/search",
        params={"format": "json", "q": q, "limit": 5, "addressdetails": 1},
        headers={"User-Agent": "EcoRevive/1.0"}
    )
    return resp.json()


@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "earth_engine": ee_initialized,
        "model": model is not None,
        "timestamp": datetime.now().isoformat()
    }


@app.post("/api/analyze", response_model=AnalyzeResponse)
async def analyze_area(request: AnalyzeRequest):
    """
    Main analysis endpoint.
    Downloads imagery, runs model, generates Gemini analysis.
    """
    try:
        # Validate bounding box
        bbox = {
            'west': request.west,
            'south': request.south,
            'east': request.east,
            'north': request.north
        }
        
        # Set date range (default to recent summer)
        start_date = request.start_date or "2023-06-01"
        end_date = request.end_date or "2023-09-30"
        
        print(f"[INFO] Analyzing area: {bbox}")
        print(f"[INFO] Date range: {start_date} to {end_date}")
        
        # Step 1: Download Sentinel-2 imagery with TILING support
        if ee_initialized:
            try:
                print("[INFO] Downloading Sentinel-2 imagery (with tiling)...")
                tiles, ee_metadata = download_sentinel2_for_model(bbox, start_date, end_date)
                print(f"   Downloaded {len(tiles)} tiles")
                
                if not tiles:
                    raise Exception("No tiles downloaded")
                    
            except Exception as e:
                print(f"[WARNING] EE download failed: {e}, using synthetic data")
                # Fallback to single synthetic tile
                synthetic_image = create_synthetic_image(bbox)
                tiles = [{'image': synthetic_image, 'row': 0, 'col': 0, 'center': (0, 0)}]
                ee_metadata = {'n_rows': 1, 'n_cols': 1}
        else:
            print("[WARNING] Earth Engine not available, using synthetic imagery for demo")
            synthetic_image = create_synthetic_image(bbox)
            tiles = [{'image': synthetic_image, 'row': 0, 'col': 0, 'center': (0, 0)}]
            ee_metadata = {'n_rows': 1, 'n_cols': 1}
        
        # Step 2: Run Fire Model inference on all tiles
        if model is None:
            raise HTTPException(status_code=503, detail="Fire model not loaded")
        
        print("[INFO] Running burn severity inference on tiles...")
        severity_map, satellite_rgb, stats = run_tiled_inference(tiles, ee_metadata)
        print(f"   Mean severity: {stats['mean_severity']:.1%}")
        print(f"   Output size: {severity_map.shape}")
        
        # Step 3: Create visualizations from stitched results
        satellite_image = rgb_array_to_base64(satellite_rgb)
        severity_image = severity_to_image(severity_map)
        raw_severity_image = severity_to_raw_image(severity_map)
        
        # Step 4: Generate Layer 2 structured output (JSON-only, single Gemini call)
        print("[INFO] Generating Layer 2 structured output...")
        layer2_output = None
        try:
            from reasoning.layer2_output import create_layer2_response
            from reasoning.gemini_multimodal import MultimodalAnalyzer
            from reasoning import create_client
            
            # Prepare RGB tile
            if satellite_rgb.ndim == 3 and satellite_rgb.shape[2] == 3:
                rgb_tile = np.moveaxis(satellite_rgb, -1, 0)
            else:
                rgb_tile = satellite_rgb
            
            center_lat = (bbox['north'] + bbox['south']) / 2
            center_lon = (bbox['west'] + bbox['east']) / 2
            location = (center_lat, center_lon)
            
            # Get structured Gemini analysis for Layer 2 (JSON-only, no verbose text)
            client = create_client()
            analyzer = MultimodalAnalyzer(client)
            l2_result = analyzer.analyze_for_layer2(
                rgb_tile=rgb_tile,
                severity_map=severity_map,
                location=location,
                unet_confidence=0.85
            )
            
            if l2_result.get('status') == 'complete':
                layer2_output = create_layer2_response(
                    severity_map=severity_map,
                    location=location,
                    bbox=bbox,
                    gemini_analysis=l2_result.get('layer2_data'),
                    model_confidence=0.85,
                    imagery_date=end_date
                )
                print("   [OK] Layer 2 JSON generated")
        except Exception as e:
            print(f"   [WARNING] Layer 2 generation failed: {e}")
            import traceback
            traceback.print_exc()

        # Step 5: Run Layer 3 contextual analysis (urban detection & cautions)
        print("[INFO] Running Layer 3 contextual analysis...")
        layer3_context = None
        try:
            from reasoning.layer3_context import create_layer3_response

            layer3_context = create_layer3_response(
                client=client,
                rgb_image=satellite_rgb,
                severity_map=severity_map,
                location=location,
                use_gemini=True
            )

            # Log the land use classification
            land_use = layer3_context.get('land_use', {})
            caution_level = layer3_context.get('overall_caution_level', 'none')
            print(f"   Land use: {land_use.get('land_use_type', 'unknown')} "
                  f"(urban: {land_use.get('urban_percentage', 0):.0f}%)")
            print(f"   Caution level: {caution_level}")
            if caution_level in ['moderate', 'high']:
                print(f"   [WARNING] {land_use.get('caution_message', '')[:100]}...")
            print("   [OK] Layer 3 context generated")
        except Exception as e:
            print(f"   [WARNING] Layer 3 analysis failed: {e}")
            import traceback
            traceback.print_exc()

        # Step 6: Calculate carbon sequestration potential
        print("[INFO] Calculating carbon sequestration potential...")
        carbon_analysis = None
        try:
            from reasoning.carbon_calculator import create_carbon_response

            # Calculate area in hectares from bbox
            lat_mid = (bbox['north'] + bbox['south']) / 2
            km_per_deg_lon = 111.32 * np.cos(lat_mid * np.pi / 180)
            km_per_deg_lat = 110.574
            width_km = abs(bbox['east'] - bbox['west']) * km_per_deg_lon
            height_km = abs(bbox['north'] - bbox['south']) * km_per_deg_lat
            area_km2 = width_km * height_km
            area_hectares = area_km2 * 100  # 1 km² = 100 hectares

            # Get land use type from Layer 3 if available
            land_use_type = None
            if layer3_context:
                land_use_type = layer3_context.get('land_use', {}).get('land_use_type')

            carbon_analysis = create_carbon_response(
                area_hectares=area_hectares,
                severity_stats=stats,
                location=location,
                user_type=request.user_type,
                land_use_type=land_use_type
            )

            # Log the results
            summary = carbon_analysis.get('summary', {})
            print(f"   Carbon potential: {summary.get('headline', 'N/A')}")
            print("   [OK] Carbon analysis complete")
        except Exception as e:
            print(f"   [WARNING] Carbon calculation failed: {e}")
            import traceback
            traceback.print_exc()

        print("[OK] Analysis complete!")

        # Fire verification: check if area is urban with no actual fire
        fire_verified = True
        fire_verification_message = None
        if layer3_context:
            land_use = layer3_context.get('land_use', {})
            land_use_type = land_use.get('land_use_type', '').lower()
            urban_pct = land_use.get('urban_percentage', 0)
            if land_use_type in ('urban', 'suburban', 'developed') or urban_pct > 50:
                # Check if severity is likely a false positive (model artifact on urban surfaces)
                high_sev = stats.get('high_severity_ratio', 0) if stats else 0
                mean_sev = stats.get('mean_severity', 0) if stats else 0
                # Urban areas with moderate-looking severity but no extreme burn = likely false positive
                if high_sev < 0.3 and mean_sev < 0.6:
                    fire_verified = False
                    fire_verification_message = (
                        f"This area is classified as {land_use_type} ({urban_pct}% developed). "
                        "No significant fire activity detected. The burn severity model is trained "
                        "on wildland fires and may produce inaccurate readings for urban areas."
                    )
                    print(f"[INFO] Urban area detected - fire_verified=False")

        return AnalyzeResponse(
            success=True,
            satellite_image=satellite_image,
            severity_image=severity_image,
            raw_severity_image=raw_severity_image,
            severity_stats=stats,
            gemini_analysis=None,  # No longer sending verbose text
            layer2_output=layer2_output,
            layer3_context=layer3_context,
            carbon_analysis=carbon_analysis,
            fire_verified=fire_verified,
            fire_verification_message=fire_verification_message
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Analysis failed: {e}")
        import traceback
        traceback.print_exc()
        return AnalyzeResponse(
            success=False,
            error=str(e)
        )


@app.post("/api/layer2-analyze", response_model=Layer2AnalyzeResponse)
async def layer2_analyze_area(request: AnalyzeRequest):
    """
    Layer 2 structured analysis endpoint.
    
    Returns structured JSON output for Layer 3 consumption.
    This endpoint is optimized for programmatic access and includes:
    - Location context (lat/lon, state, country)
    - Site characteristics (soil, terrain, land-use history)
    - Ecosystem classification
    - Computed metrics (burn %, NDVI, healing time)
    - Spatial primitives (zones, hazards, risk grid)
    - Machine-readable signals for downstream workflows
    """
    try:
        from reasoning.layer2_output import create_layer2_response
        
        # Validate bounding box
        bbox = {
            'west': request.west,
            'south': request.south,
            'east': request.east,
            'north': request.north
        }
        
        # Set date range
        start_date = request.start_date or "2023-06-01"
        end_date = request.end_date or "2023-09-30"
        
        print(f"[INFO] Layer 2 Analysis: {bbox}")
        
        # Step 1: Download Sentinel-2 imagery
        if ee_initialized:
            try:
                print("[INFO] Downloading Sentinel-2 imagery...")
                tiles, ee_metadata = download_sentinel2_for_model(bbox, start_date, end_date)
                if not tiles:
                    raise Exception("No tiles downloaded")
            except Exception as e:
                print(f"[WARNING] EE download failed: {e}, using synthetic data")
                synthetic_image = create_synthetic_image(bbox)
                tiles = [{'image': synthetic_image, 'row': 0, 'col': 0, 'center': (0, 0)}]
                ee_metadata = {'n_rows': 1, 'n_cols': 1}
        else:
            print("[WARNING] Earth Engine not available, using synthetic imagery")
            synthetic_image = create_synthetic_image(bbox)
            tiles = [{'image': synthetic_image, 'row': 0, 'col': 0, 'center': (0, 0)}]
            ee_metadata = {'n_rows': 1, 'n_cols': 1}
        
        # Step 2: Run Fire Model inference
        if model is None:
            raise HTTPException(status_code=503, detail="Fire model not loaded")
        
        print("[INFO] Running burn severity inference...")
        severity_map, satellite_rgb, stats = run_tiled_inference(tiles, ee_metadata)
        
        # Step 3: Calculate center location
        center_lat = (bbox['north'] + bbox['south']) / 2
        center_lon = (bbox['west'] + bbox['east']) / 2
        location = (center_lat, center_lon)
        
        # Step 4: Run Gemini multimodal analysis for enhancement
        gemini_analysis = None
        try:
            # Prepare RGB tile for multimodal
            if satellite_rgb.ndim == 3 and satellite_rgb.shape[2] == 3:
                rgb_tile = np.moveaxis(satellite_rgb, -1, 0)
            else:
                rgb_tile = satellite_rgb
            
            from reasoning.gemini_multimodal import MultimodalAnalyzer
            from reasoning import create_client
            
            client = create_client()
            analyzer = MultimodalAnalyzer(client)
            
            # Use analyze_for_layer2 for structured JSON output
            result = analyzer.analyze_for_layer2(
                rgb_tile=rgb_tile,
                severity_map=severity_map,
                location=location,
                unet_confidence=0.85
            )
            
            if result.get('status') == 'complete':
                # Use layer2_data which is already transformed to Layer2Output schema
                gemini_analysis = result.get('layer2_data')
                
        except Exception as e:
            print(f"[WARNING] Gemini analysis failed: {e}, proceeding without enhancement")
        
        # Step 5: Generate Layer 2 structured output
        print("[INFO] Generating Layer 2 structured output...")
        layer2_output = create_layer2_response(
            severity_map=severity_map,
            location=location,
            bbox=bbox,
            gemini_analysis=gemini_analysis,
            model_confidence=0.85,
            imagery_date=end_date
        )
        
        # Step 6: Create images for reference
        satellite_image = rgb_array_to_base64(satellite_rgb)
        severity_image = severity_to_image(severity_map)
        
        print(f"[OK] Layer 2 analysis complete! Found {len(layer2_output.get('zones', []))} zones, "
              f"{len(layer2_output.get('hazards', []))} hazards")
        
        return Layer2AnalyzeResponse(
            success=True,
            layer2_output=layer2_output,
            satellite_image=satellite_image,
            severity_image=severity_image
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"[ERROR] Layer 2 analysis failed: {e}")
        import traceback
        traceback.print_exc()
        return Layer2AnalyzeResponse(
            success=False,
            error=str(e)
        )


@app.post("/api/chat", response_model=ChatResponse)
async def chat_with_ai(request: ChatRequest):
    """
    Dynamic AI Chat endpoint.

    Uses Gemini to generate context-aware responses based on the analysis data.
    Supports quick actions (predefined prompts) and free-form questions.
    """
    try:
        from reasoning import create_client

        print(f"[INFO] Chat request: action={request.action_type}, user={request.user_type}")

        # Get the Gemini client
        client = create_client()

        # Build context from the analysis results
        context = request.context or {}
        severity_stats = context.get('severity_stats', {})
        bbox = context.get('bbox', {})
        layer2 = context.get('layer2_output', {})

        # Calculate area if bbox provided
        area_km2 = 0
        if bbox:
            lat_mid = (bbox.get('north', 0) + bbox.get('south', 0)) / 2
            km_per_deg_lon = 111.32 * np.cos(lat_mid * np.pi / 180)
            km_per_deg_lat = 110.574
            width = abs(bbox.get('east', 0) - bbox.get('west', 0)) * km_per_deg_lon
            height = abs(bbox.get('north', 0) - bbox.get('south', 0)) * km_per_deg_lat
            area_km2 = width * height

        # Build concise system context for faster responses
        high_sev = severity_stats.get('high_severity_ratio', 0) * 100
        mean_sev = severity_stats.get('mean_severity', 0) * 100
        fire_verified = context.get('fire_verified', True)
        layer3 = context.get('layer3_context', {})
        land_use = layer3.get('land_use', {}) if layer3 else {}
        land_use_type = land_use.get('land_use_type', 'unknown')

        # Location info from bbox
        lat = (bbox.get('north', 0) + bbox.get('south', 0)) / 2
        lon = (bbox.get('east', 0) + bbox.get('west', 0)) / 2

        # Retrieve RAG context from knowledge base
        rag_context = ""
        try:
            from reasoning.rag.ecology_rag import CombinedRAG
            rag = CombinedRAG()
            location_desc = f"California site at {lat:.2f}°N, {abs(lon):.2f}°W, land use: {land_use_type}"
            severity_level = "high" if high_sev > 30 else "moderate" if mean_sev > 20 else "low"
            rag_context = rag.get_full_context(
                location_description=location_desc,
                severity_level=severity_level,
                activity_type="restoration" if fire_verified else "planning"
            )
            print(f"   [RAG] Retrieved {len(rag_context)} chars of context")
        except Exception as e:
            print(f"   [WARNING] RAG context retrieval failed: {e}")

        if fire_verified is False:
            system_context = f"""Environmental advisor for EcoRevive. Location: {lat:.2f}°N, {lon:.2f}°W, {area_km2:.1f}km², land use: {land_use_type}. No recent fire activity detected here. Focus on: environmental planning, fire preparedness, urban ecology, sustainability. Be concise, use markdown."""
        else:
            system_context = f"""Restoration ecologist for EcoRevive. Site: {lat:.2f}°N, {lon:.2f}°W, {area_km2:.1f}km², land use: {land_use_type}, {mean_sev:.0f}% mean severity, {high_sev:.0f}% high severity. Be concise, use markdown."""

        # Append RAG context if available
        if rag_context:
            system_context += f"\n\n--- REFERENCE DATA ---\n{rag_context}\n--- END REFERENCE DATA ---\nUse the reference data above to ground your response with real species names, legal requirements, and ecological data. Cite specific species, laws, or permits when relevant."

        # Build short prompts for fast responses
        if request.action_type:
            action_prompts = {
                # PERSONAL - short prompts
                'safety': f"Safety checklist for volunteers at this {high_sev:.0f}% high-severity burn site. List: hazards, required gear, zones to avoid. Keep brief.",

                'hope': f"Recovery timeline for this burn site. Show: Now, Year 5, Year 10, Year 15. Include species, cover %, carbon. Be encouraging but realistic. Use species from the reference data.",

                'ownership': f"Land ownership guide for California site at {bbox.get('south', 0):.2f}°N, {abs(bbox.get('west', 0)):.2f}°W. Cover: jurisdiction, permits, contacts, timeline. Reference the legal data provided.",

                'supplies': f"Supply list & budget for 10-person restoration event at {area_km2:.1f}km² site. Table format: item, qty, cost. Include total.",

                # PROFESSIONAL - short prompts
                'legal': f"Legal/tenure analysis for professional restoration grant. Use the legal framework and land ownership data provided. Cover: ownership verification, protected status, encumbrances, compliance requirements.",

                'biophysical': f"Biophysical characterization: soil, hydrology, topography, land use history. Recommend species from the reference data for {mean_sev:.0f}% severity site.",

                'species': f"Native species palette using the species catalog provided. Tables: pioneers (0-2yr), mid-succession (2-5yr), climax (5-10yr). Include scientific names, planting density, and survival rates from the data.",

                'monitoring': f"Monitoring framework for restoration. Include: baseline metrics, schedule table (Year 0-10), carbon protocol eligibility, uncertainty bounds."
            }

            prompt = action_prompts.get(request.action_type, request.message)
        else:
            prompt = f"Question: {request.message}\nAnswer concisely based on site data and reference knowledge."

        full_prompt = f"{system_context}\n\n{prompt}"

        response = client.analyze_multimodal(
            prompt=full_prompt,
            use_json=False
        )

        print(f"   [OK] Generated response ({response['usage']['response_tokens']} tokens)")

        return ChatResponse(
            success=True,
            response=response['text']
        )

    except Exception as e:
        print(f"[ERROR] Chat failed: {e}")
        import traceback
        traceback.print_exc()
        return ChatResponse(
            success=False,
            error=str(e)
        )


@app.post("/api/export/pdf", response_model=PDFExportResponse)
async def export_pdf(request: PDFExportRequest):
    """
    Generate a PDF report for burn severity analysis.

    Supports two report types:
    - "personal": 1-2 page Impact Card (shareable, emotional)
    - "professional": 5-10 page grant-ready document (technical, comprehensive)
    """
    try:
        from reasoning.pdf_export import generate_pdf
        import base64

        print(f"[INFO] PDF Export request: type={request.report_type}, user={request.user_type}")

        # Determine report type (use user_type if report_type not explicitly set)
        report_type = request.report_type
        if report_type not in ["personal", "professional"]:
            report_type = request.user_type

        # Generate PDF
        pdf_bytes, metadata = generate_pdf(
            report_type=report_type,
            satellite_image=request.satellite_image,
            severity_image=request.severity_image,
            severity_stats=request.severity_stats,
            bbox=request.bbox,
            layer2_output=request.layer2_output,
            layer3_context=request.layer3_context,
            carbon_analysis=request.carbon_analysis,
            location_name=request.location_name,
            analysis_id=request.analysis_id,
            chat_history=request.chat_history,
        )

        # Encode to base64
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')

        print(f"   [OK] PDF generated: {metadata['filename']} ({metadata['file_size_bytes']} bytes)")

        return PDFExportResponse(
            success=True,
            pdf_base64=pdf_base64,
            filename=metadata['filename']
        )

    except Exception as e:
        print(f"[ERROR] PDF export failed: {e}")
        import traceback
        traceback.print_exc()
        return PDFExportResponse(
            success=False,
            error=str(e)
        )


@app.post("/api/export/word", response_model=WordExportResponse)
async def export_word(request: WordExportRequest):
    """
    Generate a Word document report for burn severity analysis.
    Matches PDF content so users can edit and customize.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import base64
        import io
        from PIL import Image
        from datetime import datetime

        print(f"[INFO] Word Export request: type={request.report_type}")

        # Create document
        doc = Document()

        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Calculate values
        location_name = request.location_name or "Analysis Site"
        timestamp = datetime.now().strftime("%B %d, %Y")
        area_km2 = _calculate_area_km2(request.bbox)
        area_hectares = area_km2 * 100

        # Safely extract severity values
        high_sev = float(request.severity_stats.get('high_severity_ratio', 0) or 0) * 100
        mod_sev = float(request.severity_stats.get('moderate_severity_ratio', 0) or 0) * 100
        low_sev = float(request.severity_stats.get('low_severity_ratio', 0) or 0) * 100
        mean_sev = float(request.severity_stats.get('mean_severity', 0) or 0) * 100

        # ==================== TITLE ====================
        title = doc.add_heading('EcoRevive Burn Severity Analysis', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Location header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.add_run(f"{location_name}\n").bold = True

        # Coordinates
        center_lat = (request.bbox['north'] + request.bbox['south']) / 2
        center_lon = (request.bbox['west'] + request.bbox['east']) / 2
        lat_dir = 'N' if center_lat >= 0 else 'S'
        lon_dir = 'E' if center_lon >= 0 else 'W'
        coords = f"{abs(center_lat):.4f}°{lat_dir}, {abs(center_lon):.4f}°{lon_dir}"
        header_para.add_run(f"{coords}  |  {area_km2:.2f} km² ({area_hectares:.0f} ha)  |  {timestamp}")

        doc.add_paragraph()

        # ==================== IMAGES ====================
        doc.add_heading('Satellite Analysis', level=1)

        # Add images side by side if available
        if request.satellite_image and request.severity_image:
            # Create a table for side-by-side images
            img_table = doc.add_table(rows=2, cols=2)

            # Add satellite image
            try:
                sat_data = request.satellite_image.split(',')[1] if ',' in request.satellite_image else request.satellite_image
                sat_bytes = base64.b64decode(sat_data)
                sat_stream = io.BytesIO(sat_bytes)
                cell = img_table.rows[0].cells[0]
                cell_para = cell.paragraphs[0]
                run = cell_para.add_run()
                run.add_picture(sat_stream, width=Inches(2.8))
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                img_table.rows[0].cells[0].text = "[Satellite Image]"

            # Add severity image
            try:
                sev_data = request.severity_image.split(',')[1] if ',' in request.severity_image else request.severity_image
                sev_bytes = base64.b64decode(sev_data)
                sev_stream = io.BytesIO(sev_bytes)
                cell = img_table.rows[0].cells[1]
                cell_para = cell.paragraphs[0]
                run = cell_para.add_run()
                run.add_picture(sev_stream, width=Inches(2.8))
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                img_table.rows[0].cells[1].text = "[Severity Map]"

            # Captions
            img_table.rows[1].cells[0].text = "Sentinel-2 Satellite View"
            img_table.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_table.rows[1].cells[1].text = "AI Burn Severity Map"
            img_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # ==================== EXECUTIVE SUMMARY ====================
        doc.add_heading('Executive Summary', level=1)

        summary = doc.add_paragraph()
        summary.add_run(f"This {area_km2:.1f} km² site shows ")
        if mean_sev > 60:
            run = summary.add_run("SEVERE IMPACT")
            run.bold = True
        elif mean_sev > 35:
            run = summary.add_run("MODERATE IMPACT")
            run.bold = True
        else:
            run = summary.add_run("LOW-MODERATE IMPACT")
            run.bold = True
        summary.add_run(f" with {mean_sev:.0f}% average burn severity and {high_sev:.0f}% high-severity areas.")

        doc.add_paragraph()

        # ==================== SEVERITY TABLE ====================
        doc.add_heading('Burn Severity Breakdown', level=1)

        table = doc.add_table(rows=4, cols=4)
        table.style = 'Table Grid'

        # Header row
        headers = ['Category', 'Percentage', 'Area (km²)', 'What This Means']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        severity_data = [
            ['High Severity', f'{high_sev:.1f}%', f'{high_sev * area_km2 / 100:.2f}', 'Complete vegetation loss, may need replanting'],
            ['Moderate', f'{mod_sev:.1f}%', f'{mod_sev * area_km2 / 100:.2f}', 'Partial damage, natural recovery likely'],
            ['Low/Unburned', f'{low_sev:.1f}%', f'{low_sev * area_km2 / 100:.2f}', 'Minimal impact, seed source preserved'],
        ]

        for row_idx, row_data in enumerate(severity_data):
            for col_idx, cell_data in enumerate(row_data):
                table.rows[row_idx + 1].cells[col_idx].text = cell_data

        doc.add_paragraph()

        # ==================== CARBON IMPACT ====================
        if request.carbon_analysis:
            doc.add_heading('Restoration Impact Potential', level=1)

            if request.user_type == 'personal' and request.carbon_analysis.get('personal'):
                personal = request.carbon_analysis['personal']
                total_co2 = int(personal.get('total_co2_capture_20yr', 0))

                carbon_para = doc.add_paragraph()
                run = carbon_para.add_run(f"{total_co2:,} tons CO₂")
                run.bold = True
                run.font.size = Pt(18)
                carbon_para.add_run(" captured over 20 years of restoration")
                carbon_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                equivs = personal.get('equivalencies', {})
                if equivs:
                    doc.add_paragraph()
                    doc.add_paragraph("That's equivalent to:")
                    doc.add_paragraph(f"    • {int(equivs.get('cars_off_road_for_year', 0)):,} cars removed from road for 1 year")
                    doc.add_paragraph(f"    • {int(equivs.get('tree_seedlings_grown_10yr', 0)):,} tree seedlings grown for 10 years")
                    doc.add_paragraph(f"    • {int(equivs.get('round_trip_flights_nyc_la', 0)):,} round-trip flights NYC to LA offset")
                    doc.add_paragraph(f"    • {int(equivs.get('homes_electricity_year', 0)):,} homes powered for 1 year")

            elif request.carbon_analysis.get('professional'):
                pro = request.carbon_analysis['professional']

                # Methodology badge
                method_para = doc.add_paragraph()
                run = method_para.add_run(f"Methodology: {pro.get('methodology', 'IPCC Tier 2')}")
                run.bold = True
                method_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                doc.add_paragraph()

                # Carbon Stock Summary Table
                doc.add_heading('Carbon Stock Summary', level=2)
                carbon_table = doc.add_table(rows=5, cols=3)
                carbon_table.style = 'Table Grid'

                # Headers
                carbon_headers = ['Metric', 'Value', 'Unit']
                for i, header in enumerate(carbon_headers):
                    cell = carbon_table.rows[0].cells[i]
                    cell.text = header
                    cell.paragraphs[0].runs[0].bold = True

                # Data
                carbon_data = [
                    ['Baseline Carbon Stock', f"{pro.get('baseline_carbon_tc', 0):,.1f}", 'tC'],
                    ['Carbon Lost to Fire', f"-{pro.get('carbon_lost_tc', 0):,.1f}", 'tC'],
                    ['Current Carbon Stock', f"{pro.get('current_carbon_tc', 0):,.1f}", 'tC'],
                    ['Annual Sequestration Rate', f"{pro.get('annual_sequestration_tco2e', 0):,.1f}", 'tCO₂e/year'],
                ]

                for row_idx, row_data in enumerate(carbon_data):
                    for col_idx, cell_data in enumerate(row_data):
                        carbon_table.rows[row_idx + 1].cells[col_idx].text = cell_data

                doc.add_paragraph()

                # Sequestration Projections
                doc.add_heading('Sequestration Projections', level=2)

                projections = pro.get('projections', [])
                if projections:
                    proj_table = doc.add_table(rows=len(projections) + 1, cols=3)
                    proj_table.style = 'Table Grid'

                    # Headers
                    proj_headers = ['Timeframe', 'Cumulative tCO₂e', 'Annual Rate']
                    for i, header in enumerate(proj_headers):
                        cell = proj_table.rows[0].cells[i]
                        cell.text = header
                        cell.paragraphs[0].runs[0].bold = True

                    # Data
                    for row_idx, proj in enumerate(projections):
                        proj_table.rows[row_idx + 1].cells[0].text = f"{proj.get('years', 0)} years"
                        proj_table.rows[row_idx + 1].cells[1].text = f"{proj.get('cumulative_tco2e', 0):,.0f}"
                        proj_table.rows[row_idx + 1].cells[2].text = f"{proj.get('annual_rate_tco2e', 0):,.1f} tCO₂e/yr"
                else:
                    # Generate default projections
                    annual_rate = pro.get('annual_sequestration_tco2e', 0)
                    proj_table = doc.add_table(rows=5, cols=3)
                    proj_table.style = 'Table Grid'

                    proj_headers = ['Timeframe', 'Cumulative tCO₂e', 'Annual Rate']
                    for i, header in enumerate(proj_headers):
                        cell = proj_table.rows[0].cells[i]
                        cell.text = header
                        cell.paragraphs[0].runs[0].bold = True

                    for row_idx, years in enumerate([5, 10, 15, 20]):
                        proj_table.rows[row_idx + 1].cells[0].text = f"{years} years"
                        proj_table.rows[row_idx + 1].cells[1].text = f"{(annual_rate * years):,.0f}"
                        proj_table.rows[row_idx + 1].cells[2].text = f"{annual_rate:,.1f} tCO₂e/yr"

                doc.add_paragraph()

                # Protocol Eligibility
                doc.add_heading('Carbon Credit Protocol Eligibility', level=2)

                protocols = pro.get('protocols', {})
                if protocols:
                    proto_table = doc.add_table(rows=len(protocols) + 1, cols=2)
                    proto_table.style = 'Table Grid'

                    proto_table.rows[0].cells[0].text = 'Protocol'
                    proto_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                    proto_table.rows[0].cells[1].text = 'Eligibility'
                    proto_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

                    for row_idx, (protocol, eligible) in enumerate(protocols.items()):
                        proto_name = protocol.replace('_', ' ').replace('eligible', '').strip().title()
                        proto_table.rows[row_idx + 1].cells[0].text = proto_name
                        proto_table.rows[row_idx + 1].cells[1].text = 'Eligible' if eligible else 'Not Eligible'
                else:
                    doc.add_paragraph("Protocol eligibility assessment pending. Contact EcoRevive for detailed analysis.")

                doc.add_paragraph()

                # Uncertainty Quantification
                doc.add_heading('Uncertainty Analysis', level=2)

                ci_low = pro.get('confidence_interval_low', 0)
                ci_high = pro.get('confidence_interval_high', 0)
                uncertainty = pro.get('uncertainty_pct', 25)

                uncert_para = doc.add_paragraph()
                uncert_para.add_run("95% Confidence Interval (20-year projection): ").bold = True
                uncert_para.add_run(f"{ci_low:,.0f} - {ci_high:,.0f} tCO₂e")
                doc.add_paragraph(f"Combined Uncertainty: ±{uncertainty}%")

                doc.add_paragraph()

                # Limitations
                doc.add_heading('Methodological Limitations', level=2)

                limitations = pro.get('limitations', [
                    'Remote sensing estimates require ground-truth verification',
                    'Carbon stock calculations assume typical forest composition',
                    'Sequestration rates may vary based on management practices',
                    'Does not account for natural disturbance risk',
                    'Baseline estimates derived from regional averages'
                ])

                for lim in limitations:
                    doc.add_paragraph(f"• {lim}")

                doc.add_paragraph()

                # Data Sources
                doc.add_heading('Data Sources', level=2)

                sources = pro.get('data_sources', [
                    'Sentinel-2 MSI multispectral imagery (10m resolution)',
                    'U-Net deep learning model trained on California wildfires',
                    'IPCC Guidelines for National GHG Inventories (2006, 2019 Refinement)',
                    'FIA forest inventory regional parameters',
                    'Gemini multimodal AI for contextual analysis'
                ])

                for src in sources:
                    doc.add_paragraph(f"• {src}")

        doc.add_paragraph()

        # ==================== SITE CONTEXT ====================
        if request.layer3_context and request.layer3_context.get('land_use'):
            doc.add_heading('Site Context', level=1)
            land_use = request.layer3_context.get('land_use', {})
            land_type = land_use.get('land_use_type', 'Unknown').title()
            doc.add_paragraph(f"Land Classification: {land_type}")
            if land_use.get('land_use_description'):
                doc.add_paragraph(land_use.get('land_use_description'))

        # ==================== SAFETY NOTICE ====================
        doc.add_heading('Safety Notice', level=1)

        if high_sev > 40:
            level = "HIGH"
            message = "This site has significant high-severity burn areas. Before visiting, be aware of hazards like standing dead trees (widowmakers), unstable slopes, and ash pits. Always go with a buddy and inform someone of your plans."
        elif high_sev > 20:
            level = "MODERATE"
            message = "Exercise caution when visiting this site. Some areas may have standing dead trees and loose soil. Wear sturdy boots and bring plenty of water."
        else:
            level = "LOW"
            message = "This site appears relatively safe for visits, but always be aware of your surroundings. Standard outdoor safety precautions apply."

        safety_para = doc.add_paragraph()
        run = safety_para.add_run(f"SAFETY LEVEL: {level}")
        run.bold = True
        doc.add_paragraph(message)

        doc.add_paragraph()

        # ==================== HOW YOU CAN HELP ====================
        doc.add_heading('How You Can Help', level=1)

        actions = [
            ("Learn About Native Species", "Research which native plants thrive in your region. Focus on fire-adapted species."),
            ("Join Restoration Events", "Connect with local conservation groups that organize volunteer planting days."),
            ("Support Organizations", "Donate to or volunteer with groups like One Tree Planted or local land trusts."),
            ("Monitor & Document", "If you visit the site, document recovery with photos for citizen science."),
            ("Reduce Fire Risk", "Create defensible space around homes and practice fire-safe behaviors."),
            ("Spread Awareness", "Share information about wildfire recovery with your community."),
        ]

        for i, (title_text, desc) in enumerate(actions, 1):
            action_para = doc.add_paragraph()
            action_para.add_run(f"{i}. {title_text}: ").bold = True
            action_para.add_run(desc)

        doc.add_paragraph()

        # ==================== RECOVERY TIMELINE ====================
        doc.add_heading('Recovery Timeline', level=1)

        timeline_table = doc.add_table(rows=5, cols=2)
        timeline_table.style = 'Table Grid'

        timeline_data = [
            ['Timeframe', 'What to Expect'],
            ['Year 1-2', 'Ground cover begins returning. Pioneer grasses and shrubs establish.'],
            ['Year 3-5', 'Shrub layer develops. Tree seedlings visible. Wildlife returns.'],
            ['Year 5-10', 'Young forest structure emerges. Canopy begins closing.'],
            ['Year 10-20', 'Maturing forest. Carbon sequestration accelerates.'],
        ]

        for row_idx, row_data in enumerate(timeline_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = timeline_table.rows[row_idx].cells[col_idx]
                cell.text = cell_data
                if row_idx == 0:
                    cell.paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # ==================== CONSULTATION LOG (if chat history provided) ====================
        if request.chat_history and len(request.chat_history) > 0:
            doc.add_page_break()
            doc.add_heading('Appendix: AI Consultation Log', level=1)
            
            intro = doc.add_paragraph()
            intro.add_run(
                "The following is a record of questions asked and responses received during the analysis session. "
                "This log is provided for transparency and reference."
            )
            intro.paragraph_format.space_after = Pt(12)
            
            # Limit to last 20 messages to prevent extremely long reports
            messages_to_render = request.chat_history[-20:] if len(request.chat_history) > 20 else request.chat_history
            
            if len(request.chat_history) > 20:
                truncation_note = doc.add_paragraph()
                truncation_note.add_run(f"(Showing last 20 of {len(request.chat_history)} messages)")
                truncation_note.runs[0].italic = True
                truncation_note.paragraph_format.space_after = Pt(6)
            
            for idx, msg in enumerate(messages_to_render):
                role = msg.get('role', 'unknown').capitalize()
                content = msg.get('content', '')
                msg_timestamp = msg.get('timestamp', '')
                
                # Create message header
                msg_header = doc.add_paragraph()
                role_run = msg_header.add_run(f"{role}")
                role_run.bold = True
                if role == 'User':
                    role_run.font.color.rgb = RGBColor(0, 100, 150)  # Blue for user
                else:
                    role_run.font.color.rgb = RGBColor(0, 120, 80)  # Green for assistant
                
                if msg_timestamp:
                    # Format timestamp for display
                    try:
                        from datetime import datetime as dt
                        ts = dt.fromisoformat(msg_timestamp.replace('Z', '+00:00'))
                        ts_str = ts.strftime('%I:%M %p')
                        msg_header.add_run(f" • {ts_str}")
                    except:
                        pass
                
                msg_header.paragraph_format.space_before = Pt(8)
                msg_header.paragraph_format.space_after = Pt(2)
                
                # Create message content
                msg_body = doc.add_paragraph(content)
                msg_body.paragraph_format.left_indent = Inches(0.25)
                msg_body.paragraph_format.space_after = Pt(6)
            
            doc.add_paragraph()

        # ==================== FOOTER ====================
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(f"Generated by EcoRevive on {timestamp}").italic = True
        footer.add_run("\nData sources: Sentinel-2 satellite imagery, U-Net deep learning model, Gemini AI analysis")
        footer.add_run("\nThis report is for informational purposes. Verify with ground-truth data before making decisions.")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Save to bytes
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        docx_bytes = docx_buffer.read()

        # Encode to base64
        docx_base64 = base64.b64encode(docx_bytes).decode('utf-8')

        filename = f"EcoRevive_{location_name.replace(' ', '_').replace(',', '')}_{datetime.now().strftime('%Y%m%d')}.docx"

        print(f"   [OK] Word document generated: {filename}")

        return WordExportResponse(
            success=True,
            docx_base64=docx_base64,
            filename=filename
        )

    except ImportError:
        return WordExportResponse(
            success=False,
            error="python-docx library not installed. Install with: pip install python-docx"
        )
    except Exception as e:
        print(f"[ERROR] Word export failed: {e}")
        import traceback
        traceback.print_exc()
        return WordExportResponse(
            success=False,
            error=str(e)
        )


def _calculate_area_km2(bbox: Dict[str, float]) -> float:
    """Calculate area in km² from bbox."""
    import math
    lat_mid = (bbox['north'] + bbox['south']) / 2
    km_per_deg_lat = 110.574
    km_per_deg_lon = 111.32 * math.cos(math.radians(lat_mid))
    width = abs(bbox['east'] - bbox['west']) * km_per_deg_lon
    height = abs(bbox['north'] - bbox['south']) * km_per_deg_lat
    return width * height


@app.post("/api/hope-visualization", response_model=HopeVisualizationResponse)
async def generate_hope_visualization(request: HopeVisualizationRequest):
    """
    Generate hope visualization using Imagen 3.

    Creates AI-generated photorealistic images showing ecosystem recovery
    at different time points (5, 10, 15 years).
    """
    try:
        from reasoning.gemini_hope import HopeVisualizer
        from reasoning import create_client
        import io

        print(f"[INFO] Hope visualization request: {request.ecosystem_type}, {request.years_in_future} years")

        # Create Gemini client and visualizer
        client = create_client()
        visualizer = HopeVisualizer(gemini_client=client)

        # Generate recovery forecast first (text-based timeline)
        forecast = visualizer.forecast_recovery(
            ecosystem_type=request.ecosystem_type,
            mean_severity=request.mean_severity,
            restoration_method="combination",
            area_hectares=request.area_hectares
        )

        # Try to generate image with Imagen 3
        image_result = visualizer.generate_hope_image(
            ecosystem_type=request.ecosystem_type,
            years_in_future=request.years_in_future,
        )

        image_base64 = None
        description = None

        # Check if we got an image
        if image_result.get("image"):
            # Convert PIL image to base64
            img = image_result["image"]
            buffer = io.BytesIO()
            img.save(buffer, format="PNG")
            image_base64 = f"data:image/png;base64,{base64.b64encode(buffer.getvalue()).decode('utf-8')}"
            print(f"   [OK] Generated Imagen 3 visualization")
        elif image_result.get("fallback_description"):
            # Use text description as fallback
            description = image_result["fallback_description"]
            print(f"   [WARNING] Using text fallback (Imagen unavailable)")

        # Extract hope message from forecast
        hope_message = forecast.get("hope_message", "")
        if not description and hope_message:
            description = hope_message

        return HopeVisualizationResponse(
            success=True,
            image_base64=image_base64,
            forecast=forecast,
            description=description
        )

    except Exception as e:
        print(f"[ERROR] Hope visualization failed: {e}")
        import traceback
        traceback.print_exc()
        return HopeVisualizationResponse(
            success=False,
            error=str(e)
        )


# Run with: uvicorn server:app --reload --port 8000
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

